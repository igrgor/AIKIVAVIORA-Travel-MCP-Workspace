# -*- coding: utf-8 -*-
"""Сборка docx для ДЗ-1 Школа 6 — Travel / Hotel Analytics Pro (ПЗ + отчёт).

Запуск из корня travel workspace:
  python exports/build_docx_travel_hotel_dz1.py

Скрины: извлекаются из существующего отчёта в папке сдачи (G:),
        либо из screenshots/ если положить вручную.
"""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
EXPORTS = ROOT / "exports"
SCREENSHOTS = ROOT / "screenshots"

REPORT_NAMES = [
    "01_ui_three_columns.png",
    "02_kpi_hotels.png",
    "03_weather_mcp.png",
    "04_comparison_matrix.png",
    "05_watchlist_reports.png",
    "06_github_repo.png",
]

REPORT_CAPTIONS = [
    "Три колонки: чат, аналитика, журнал активности",
    "KPI, каталог отелей, вкладка «Все (847)»",
    "Погода в регионе отеля через Weather MCP",
    "Матрица сравнения отелей",
    "Мой список и сохранённые отчёты",
    "GitHub: AIKIVAVIORA-Travel-MCP-Workspace",
]


def find_hotels_folder() -> Path:
    root = Path("G:/19_Университет AI")
    for d in root.rglob("*"):
        if d.is_dir() and "отели" in d.name.lower() and "готово" in d.name.lower():
            return d
    raise FileNotFoundError("Папка ДЗ-1 отели не найдена на G:")


def hotel_docx_pair(folder: Path) -> tuple[Path, Path]:
    candidates = [
        f
        for f in folder.glob("*.docx")
        if ".bak" not in f.name and "Hotel Analytics" in f.name
    ]
    if len(candidates) < 2:
        raise FileNotFoundError("Не найдены ПЗ и отчёт Hotel Analytics в папке сдачи")
    report = max(candidates, key=lambda p: p.stat().st_size)
    pz = min(candidates, key=lambda p: p.stat().st_size)
    return report, pz


def setup_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("AIKIVAVIORA")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Travel MCP · Hotel Analytics Pro · Школа 6 · День 1 · Июнь 2026"
    ).italic = True

    doc.add_paragraph()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def meta_lines(doc: Document, rows: list[tuple[str, str]]) -> None:
    for label, val in rows:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(val)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image_if_exists(doc: Document, path: Path, caption: str) -> bool:
    if not path.is_file():
        p = doc.add_paragraph()
        p.add_run(f"[Скрин: {path.name}] ").italic = True
        p.add_run(caption)
        return False
    doc.add_paragraph(caption, style="Intense Quote")
    doc.add_picture(str(path), width=Cm(15.5))
    doc.add_paragraph()
    return True


def extract_images_from_docx(src: Path, dest_dir: Path) -> list[Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    extracted: list[Path] = []
    with zipfile.ZipFile(src, "r") as zf:
        media = sorted(
            n for n in zf.namelist() if n.startswith("word/media/") and "." in n
        )
        for i, name in enumerate(media):
            data = zf.read(name)
            ext = Path(name).suffix or ".png"
            out = dest_dir / f"legacy_{i + 1:02d}{ext}"
            out.write_bytes(data)
            extracted.append(out)
    return extracted


def map_screenshots(legacy: list[Path]) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for i, logical in enumerate(REPORT_NAMES):
        manual = SCREENSHOTS / logical
        if manual.is_file():
            mapping[logical] = manual
        elif i < len(legacy):
            mapping[logical] = legacy[i]
    return mapping


def build_pz() -> Path:
    doc = setup_doc("Пояснительная записка")
    meta_lines(
        doc,
        [
            ("Студент:", "Игорь Кашинцев"),
            ("Курс:", "Школа 6 — MCP для AI-агентов, день #1"),
            ("Проект:", "AIKIVAVIORA Travel MCP Workspace — Hotel Analytics Pro"),
            ("Дата:", "18.06.2026"),
            ("GitHub:", "https://github.com/igrgor/AIKIVAVIORA-Travel-MCP-Workspace"),
            ("Версия:", "v1.1 — Replit UI Shell + доработки Cursor"),
        ],
    )

    doc.add_heading("1. Назначение", level=1)
    doc.add_paragraph(
        "Учебный research-workspace для отельной аналитики (часть 2 ДЗ-1): "
        "три колонки, демо-каталог, сравнение, отчёты, Weather MCP."
    )

    doc.add_heading("2. Архитектура", level=1)
    add_bullets(
        doc,
        [
            "React UI (hotel-analytics) + Express API (api-server)",
            "Weather MCP: @dangahagan/weather-mcp через stdio",
            "localStorage: мой список, сравнение, отчёты",
            "Демо 24 отеля; целевой каталог STR 847 — в разработке",
        ],
    )

    doc.add_heading("3. Реализовано", level=1)
    add_table(
        doc,
        ["Модуль", "Статус"],
        [
            ("3-панельный UI", "✅"),
            ("Мой список (watchlist)", "✅"),
            ("Сравнение до 9 отелей", "✅"),
            ("Сохранённые отчёты (MD)", "✅"),
            ("Weather MCP + график", "✅"),
            ("Метки Демо / В разработке", "✅"),
            ("STR 847 полный каталог", "В разработке"),
            ("DeepSeek live chat", "В разработке"),
            ("PDF export", "В разработке"),
        ],
    )

    doc.add_heading("4. Ограничения", level=1)
    doc.add_paragraph(
        "KPI, журнал и источники данных — демо. Это отмечено в интерфейсе "
        "бейджами, чтобы не воспринимать как сбой. Replit-архив: Hotel-Insight-Hub.zip."
    )

    out = EXPORTS / "PZ_travel-hotel-analytics_пояснительная_записка.docx"
    EXPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_report(screens: dict[str, Path]) -> Path:
    doc = setup_doc("ОТЧЁТ О ВЫПОЛНЕННОЙ РАБОТЕ")
    meta_lines(
        doc,
        [
            ("Студент:", "Игорь Кашинцев"),
            ("Курс:", "Школа 6 — MCP для AI-агентов, день #1"),
            ("Проект:", "Hotel Analytics Pro (Travel MCP)"),
            ("Дата:", "18.06.2026"),
            ("GitHub:", "https://github.com/igrgor/AIKIVAVIORA-Travel-MCP-Workspace"),
            ("Часть ДЗ:", "2 — отели (вариант «под себя»)"),
        ],
    )

    doc.add_heading("1. Цель работы", level=1)
    doc.add_paragraph(
        "Создать учебный MCP-workspace для отельной аналитики в стиле AIKIVAVIORA: "
        "мониторинг объектов, сравнение, исследовательский чат, погода через MCP, "
        "сохранение отчётов."
    )

    doc.add_heading("2. Реализованный функционал", level=1)
    add_table(
        doc,
        ["Блок", "Описание"],
        [
            ("UI Shell (Replit)", "Три колонки, KPI, карточки отелей, матрица"),
            ("Watchlist", "Мой список, localStorage, вкладка «Мои»"),
            ("Сравнение", "До 9 отелей, снимок в отчёты"),
            ("Отчёты", "Из чата и сравнения, экспорт Markdown"),
            ("Weather MCP", "Текущая погода, 7 дней, почасово, график 30 дней"),
            ("UX прозрачность", "Бейджи: Демо, В разработке, Заглушка, Локально"),
        ],
    )

    doc.add_heading("3. Тест-сценарии", level=1)

    doc.add_heading("Сценарий 1 — Интерфейс и каталог", level=2)
    add_table(
        doc,
        ["Поле", "Результат"],
        [
            ("Действие", "Статус-бар 24/847 → вкладка «Все (847)»"),
            ("Ожидание", "Баннер «в разработке», 24 демо-отеля"),
            ("Скрин", REPORT_NAMES[1]),
        ],
    )
    add_image_if_exists(doc, screens.get(REPORT_NAMES[1], Path()), REPORT_CAPTIONS[1])

    doc.add_heading("Сценарий 2 — Weather MCP", level=2)
    add_table(
        doc,
        ["Поле", "Результат"],
        [
            ("Действие", "Выбрать отель → «Погода в регионе»"),
            ("API", "GET /api/weather/* → Weather MCP"),
            ("Скрин", REPORT_NAMES[2]),
        ],
    )
    add_image_if_exists(doc, screens.get(REPORT_NAMES[2], Path()), REPORT_CAPTIONS[2])

    doc.add_heading("Сценарий 3 — Сравнение и отчёты", level=2)
    add_table(
        doc,
        ["Поле", "Результат"],
        [
            ("Действие", "Сравнить отели → сохранить снимок; чат → сохранить отчёт"),
            ("Хранение", "localStorage (бейдж «Локально»)"),
            ("Скрины", f"{REPORT_NAMES[3]}, {REPORT_NAMES[4]}"),
        ],
    )
    for key, cap in zip(REPORT_NAMES[3:5], REPORT_CAPTIONS[3:5]):
        add_image_if_exists(doc, screens.get(key, Path()), cap)

    doc.add_heading("4. Скрины процесса", level=1)
    for key, cap in zip([REPORT_NAMES[0], REPORT_NAMES[5]], [REPORT_CAPTIONS[0], REPORT_CAPTIONS[5]]):
        add_image_if_exists(doc, screens.get(key, Path()), cap)

    doc.add_heading("5. В разработке (не баг)", level=1)
    add_bullets(
        doc,
        [
            "Полный каталог STR 847",
            "DeepSeek / OpenRouter — живой чат",
            "STR, HotStats, OTA Insight — реальные коннекторы",
            "Postgres + синхронизация отчётов",
            "Экспорт PDF",
        ],
    )

    doc.add_heading("6. Чеклист сдачи", level=1)
    add_table(
        doc,
        ["Пункт", "Статус"],
        [
            ("GitHub / архив", "✅"),
            ("2–3 теста", "✅"),
            ("Скриншоты", "✅"),
            ("ПЗ", "✅"),
            ("Бланк AIKIVAVIORA", "✅"),
        ],
    )

    out = EXPORTS / "PZ_travel-hotel-analytics_отчёт_для_сдачи.docx"
    doc.save(out)
    return out


def backup(path: Path) -> None:
    bak = path.with_suffix(path.suffix + ".bak")
    if not bak.exists() and path.exists():
        shutil.copy2(path, bak)
        print("backup:", bak.name)


def main() -> None:
    folder = find_hotels_folder()
    report_dst, pz_dst = hotel_docx_pair(folder)

    backup(report_dst)
    backup(pz_dst)

    legacy_dir = EXPORTS / "screenshots_legacy"
    legacy = extract_images_from_docx(report_dst, legacy_dir)
    print(f"extracted {len(legacy)} images from", report_dst.name)

    screens = map_screenshots(legacy)

    pz_out = build_pz()
    report_out = build_report(screens)

    shutil.copy2(pz_out, pz_dst)
    shutil.copy2(report_out, report_dst)

    for md_name in (
        "PZ_travel-hotel-analytics_пояснительная_записка.md",
        "PZ_travel-hotel-analytics_отчёт_для_сдачи.md",
    ):
        src = DOCS / md_name
        if src.is_file():
            shutil.copy2(src, folder / md_name)

    shots_dir = folder / "screenshots"
    shots_dir.mkdir(exist_ok=True)
    if SCREENSHOTS.is_dir():
        for name in REPORT_NAMES:
            src = SCREENSHOTS / name
            if src.is_file():
                shutil.copy2(src, shots_dir / name)
                shutil.copy2(src, folder / name)

    from docx import Document as D

    for path in (report_dst, pz_dst):
        imgs = sum(1 for rel in D(path).part.rels.values() if "image" in rel.reltype)
        print(f"OK: {path.name} | {path.stat().st_size // 1024} KB | images: {imgs}")

    print("folder:", folder)
    print("exports:", EXPORTS)


if __name__ == "__main__":
    main()
