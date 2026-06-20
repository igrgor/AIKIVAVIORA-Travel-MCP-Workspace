# -*- coding: utf-8 -*-
"""Удалить секреты (ghp_*) из ТЗ docx в папке сдачи."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

TOKEN_RE = re.compile(r"ghp_[A-Za-z0-9_]+")


def find_hotels_folder() -> Path:
    root = Path("G:/19_Университет AI")
    for d in root.rglob("*"):
        if d.is_dir() and "отели" in d.name.lower() and "готово" in d.name.lower():
            return d
    raise FileNotFoundError("Папка сдачи не найдена")


def redact_docx(path: Path) -> int:
    doc = Document(path)
    count = 0
    for para in doc.paragraphs:
        if TOKEN_RE.search(para.text):
            for run in para.runs:
                if TOKEN_RE.search(run.text):
                    run.text = TOKEN_RE.sub("[ТОКЕН УДАЛЁН — отзовите в GitHub]", run.text)
                    count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if TOKEN_RE.search(run.text):
                            run.text = TOKEN_RE.sub(
                                "[ТОКЕН УДАЛЁН — отзовите в GitHub]", run.text
                            )
                            count += 1
    if count:
        doc.save(path)
    return count


def main() -> None:
    folder = find_hotels_folder()
    total = 0
    for path in folder.glob("*.docx"):
        if ".bak" in path.name or "Hotel Analytics" in path.name:
            continue
        n = redact_docx(path)
        if n:
            print(f"redacted {n} runs in {path.name}")
            total += n
    if total == 0:
        print("no tokens found")
    else:
        print("total redactions:", total)


if __name__ == "__main__":
    main()
